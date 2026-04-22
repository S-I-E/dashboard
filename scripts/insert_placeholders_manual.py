#!/usr/bin/env python3
# Best-effort inserter de placeholders em um DOCX específico usando python-docx
import sys
from pathlib import Path

TEMPLATES_DIR = Path.cwd() / 'templates' / 'legal_instruments' / 'generated'

LABELS = [
    'Cidade', 'UF', 'CEP', 'Representante legal', 'C.P.F./M.F.', 'CPF', 'CNPJ', 'Identidade n.º', 'Nacionalidade', 'Cargo', 'Ato de Nomeação',
    'Instituição', 'Natureza Jurídica', 'Endereço', 'Objeto', 'Plano de trabalho', 'Vigência', 'Valor', 'Prazo', 'Prazo de vigência', 'Data', 'Partes',
]

# normalize label to id
import re
import unicodedata

def slugify(s):
    s = unicodedata.normalize('NFKD', s).encode('ascii','ignore').decode('ascii')
    s = re.sub(r'[^a-zA-Z0-9]+','_', s)
    return s.strip('_').lower()

try:
    import docx
except Exception:
    print('python-docx not installed, attempting to install...')
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx


def process_file(filename: str):
    path = TEMPLATES_DIR / filename
    if not path.exists():
        print('File not found:', path)
        return
    doc = docx.Document(path)
    inserted = set()
    for p in doc.paragraphs:
        text = p.text
        for label in LABELS:
            # look for 'Label:' or 'Label -' or 'Label\t'
            m = re.search(r'(' + re.escape(label) + r')\s*[:\-\t]\s*(.+)$', text)
            if m:
                # replace the trailing content with placeholder
                id = slugify(label)
                placeholder = '{{' + id + '}}'
                # modify runs: remove runs after label and append placeholder run
                # simple approach: clear paragraph and rebuild
                new_text = text[:m.start(2)] + placeholder
                # set paragraph text
                for r in list(p.runs):
                    r.text = ''
                p.add_run(new_text)
                inserted.add(id)
                break
    # save to clean folder
    out_dir = TEMPLATES_DIR / 'clean'
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / path.name
    doc.save(out_path)
    # generate json fields
    fields = []
    for id in sorted(inserted):
        fields.append({'id': id, 'name': id, 'type': 'text', 'label': id.replace('_',' ').capitalize(), 'required': False})
    import json
    json_path = out_dir / (path.stem + '.json')
    json_obj = {'template': out_path.name, 'fields': fields}
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_obj, f, ensure_ascii=False, indent=2)
    print('Wrote edited docx to', out_path)
    print('Wrote json to', json_path)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: insert_placeholders_manual.py <filename.docx>')
        sys.exit(2)
    process_file(sys.argv[1])
